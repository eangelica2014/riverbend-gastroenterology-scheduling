from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "artifacts" / "Riverbend_Gastroenterology_Scheduling_PRD_Eangelica_Aton.docx"
LOGO = Path("/home/ubuntu/webdev-static-assets/riverbend-logo.png")
ARCH = ROOT / "artifacts" / "diagrams" / "runtime-architecture.png"
LIFECYCLE = ROOT / "artifacts" / "diagrams" / "fieldflow-lifecycle.png"
METHOD = ROOT / "artifacts" / "diagrams" / "dual-method.png"
SCENARIOS = ROOT / "agent" / "scenarios.json"

INK = "102A24"
INK_SOFT = "26463E"
PAPER = "F4F0E6"
PAPER_LIGHT = "FBF8F0"
PAPER_DEEP = "E9E1D2"
PERSIMMON = "F05A28"
CHARTREUSE = "B9D531"
GRAPHITE = "6B7069"
RED = "B42318"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text_color(cell, color: str, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)
            run.bold = bold


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PAGE ")
    run.font.name = "IBM Plex Mono"
    run.font.size = Pt(8)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_sep])
    text = paragraph.add_run("Right-click this field and select Update Field if the table of contents is not populated.")
    text.font.color.rgb = RGBColor.from_string(GRAPHITE)
    text.italic = True
    run._r.append(fld_end)


def configure_section(section, landscape: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)
    section.top_margin = Inches(0.64)
    section.bottom_margin = Inches(0.62)
    section.header_distance = Inches(0.26)
    section.footer_distance = Inches(0.28)

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(0.28))
    run = p.add_run("  RIVERBEND / CONFIDO PRODUCT TAKE-HOME   ·   EANGELICA ATON")
    run.font.name = "IBM Plex Mono"
    run.font.size = Pt(7)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(INK_SOFT)

    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=section.page_width - section.left_margin - section.right_margin)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    left = table.cell(0, 0).paragraphs[0]
    left_run = left.add_run("SYNTHETIC DATA ONLY   ·   POLICY v1.0.0   ·   26 AUG 2026")
    left_run.font.name = "IBM Plex Mono"
    left_run.font.size = Pt(7)
    left_run.font.color.rgb = RGBColor.from_string(PERSIMMON)
    add_page_number(table.cell(0, 1).paragraphs[0])


def apply_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "IBM Plex Sans"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in (("Title", 44, INK), ("Heading 1", 25, INK), ("Heading 2", 17, INK), ("Heading 3", 11.5, PERSIMMON)):
        style = styles[name]
        style.font.name = "Newsreader"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Title"
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(11 if name != "Heading 1" else 18)
        style.paragraph_format.space_after = Pt(6)

    styles["Heading 1"].paragraph_format.page_break_before = False
    styles["Heading 3"].font.name = "IBM Plex Mono"
    styles["Heading 3"].font.all_caps = True
    styles["Heading 3"].font.size = Pt(9)

    quote = styles["Quote"]
    quote.font.name = "Newsreader"
    quote.font.size = Pt(14)
    quote.font.italic = True
    quote.font.color.rgb = RGBColor.from_string(INK)
    quote.paragraph_format.left_indent = Inches(0.28)
    quote.paragraph_format.right_indent = Inches(0.18)
    quote.paragraph_format.space_before = Pt(8)
    quote.paragraph_format.space_after = Pt(10)


def add_heading(doc: Document, text: str, level: int = 1, kicker: str | None = None) -> None:
    if kicker:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        if level == 1:
            p.paragraph_format.page_break_before = True
        run = p.add_run(kicker.upper())
        run.font.name = "IBM Plex Mono"
        run.font.size = Pt(7.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(PERSIMMON)
        run.font.letter_spacing = Pt(1) if hasattr(run.font, "letter_spacing") else None
    heading = doc.add_heading(text, level=level)
    if level == 1 and not kicker:
        heading.paragraph_format.page_break_before = True


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first, rest = text.split(bold_prefix, 1)
        p.add_run(first)
        p.add_run(bold_prefix).bold = True
        p.add_run(rest)
    else:
        p.add_run(text)


def add_callout(doc: Document, title: str, text: str, tone: str = "ink") -> None:
    fill = {"ink": INK, "attention": PERSIMMON, "verified": CHARTREUSE, "paper": PAPER_DEEP}[tone]
    text_color = WHITE if tone in {"ink", "attention"} else INK
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, start=170, bottom=140, end=170)
    p = cell.paragraphs[0]
    r1 = p.add_run(title.upper() + "\n")
    r1.font.name = "IBM Plex Mono"
    r1.font.size = Pt(8)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor.from_string(text_color if tone != "verified" else INK)
    r2 = p.add_run(text)
    r2.font.name = "Newsreader"
    r2.font.size = Pt(15)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor.from_string(text_color)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]], widths: Sequence[float] | None = None, font_size: float = 7.6) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = widths is None
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, INK)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(text).upper())
        run.font.name = "IBM Plex Mono"
        run.font.size = Pt(font_size)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(CHARTREUSE)
        if widths:
            cell.width = Inches(widths[index])

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_index % 2:
                set_cell_shading(cell, PAPER_DEEP)
            else:
                set_cell_shading(cell, PAPER_LIGHT)
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.name = "IBM Plex Sans"
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor.from_string(INK)
            if widths:
                cell.width = Inches(widths[index])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, path: Path, caption: str, width: float = 7.0) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cp.add_run(caption)
    run.font.name = "IBM Plex Mono"
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor.from_string(GRAPHITE)
    run.italic = True


def add_bullets(doc: Document, items: Sequence[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.add_run(item)


def add_reference(doc: Document, number: int, title: str, url: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(f"[{number}] {title}. ")
    r.bold = True
    u = p.add_run(url)
    u.font.color.rgb = RGBColor(5, 99, 193)
    u.underline = True


def build_document() -> None:
    doc = Document()
    doc.core_properties.title = "Riverbend Gastroenterology Scheduling — Product Requirements and Platform Strategy"
    doc.core_properties.subject = "Confido Health Product Manager, Platform take-home"
    doc.core_properties.author = "Eangelica Germano Aton"
    doc.core_properties.keywords = "Confido Health, voice agent, healthcare scheduling, FDE platform, FieldFlow, PRD"
    apply_styles(doc)
    configure_section(doc.sections[0])

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(38)
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(0.9))
    badge = doc.add_table(rows=1, cols=1)
    badge.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = badge.cell(0, 0)
    set_cell_shading(cell, CHARTREUSE)
    set_cell_margins(cell, top=50, start=90, bottom=50, end=90)
    br = cell.paragraphs[0].add_run("CONFIDO HEALTH · PRODUCT MANAGER, PLATFORM TAKE-HOME")
    br.font.name = "IBM Plex Mono"
    br.font.size = Pt(8)
    br.font.bold = True
    br.font.color.rgb = RGBColor.from_string(INK)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("Riverbend\nGastroenterology\nScheduling")
    r.font.name = "Newsreader"
    r.font.size = Pt(42)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(INK)

    subtitle = doc.add_paragraph()
    sr = subtitle.add_run("Voice Agent PRD · FieldFlow Platform Strategy · Experiment Portfolio")
    sr.font.name = "IBM Plex Sans"
    sr.font.size = Pt(13)
    sr.font.color.rgb = RGBColor.from_string(PERSIMMON)
    subtitle.paragraph_format.space_after = Pt(22)

    add_callout(doc, "Operating thesis", "The conversation can be flexible. The policy cannot.", "ink")

    byline = doc.add_paragraph()
    byline.paragraph_format.space_before = Pt(20)
    rb = byline.add_run("Eangelica Germano Aton\n")
    rb.font.name = "Newsreader"
    rb.font.size = Pt(17)
    rb.font.bold = True
    rm = byline.add_run("Technical Product Lead · NYC · August 26, 2026")
    rm.font.name = "IBM Plex Mono"
    rm.font.size = Pt(8)
    rm.font.color.rgb = RGBColor.from_string(GRAPHITE)

    notice = doc.add_paragraph()
    notice.paragraph_format.space_before = Pt(18)
    nr = notice.add_run("SYNTHETIC DATA ONLY  ·  NO LIVE EHR  ·  NO PRODUCTION TELEPHONY")
    nr.font.name = "IBM Plex Mono"
    nr.font.size = Pt(8)
    nr.font.bold = True
    nr.font.color.rgb = RGBColor.from_string(RED)

    # TOC
    add_heading(doc, "Document Guide", 1, "How to use this PRD")
    add_body(doc, "This native Microsoft Word document summarizes the complete Phase 1 and Phase 2 solution. The interactive prototype and source repository remain the executable evidence; this PRD explains the decisions, contracts, requirements, tests, platform thesis, experiments, and candidate operating method behind them.")
    toc = doc.add_paragraph()
    add_toc(toc)
    add_callout(doc, "Evaluator path", "Read the executive decision, run the agent and scenario lab, then inspect FieldFlow and the P0 experiments.", "verified")

    # Executive decision
    add_heading(doc, "Executive Decision", 1, "Summary")
    add_body(doc, "Riverbend is a bounded administrative voice agent for appointment management, clinic information, and deliberate human transfers. Flexible conversation is separated from deterministic authority: safety and identity gates run before protected work; a versioned policy engine evaluates eligibility and provider rules before availability; typed tools mediate every read, write, and transfer; and a trace preserves the reason, rule, tool, and result for inspection.")
    add_body(doc, "Phase 1 demonstrates this system with synthetic patient, provider, slot, and appointment fixtures; browser voice input/output; a fully operable text path; 24 replayable scenarios; and 18 automated assertions. The suite, TypeScript check, and production build pass. The prototype makes no claim of production telephony, live EHR integration, or clinical decision support.")
    add_body(doc, "Phase 2 proposes FieldFlow, an FDE-first control plane that turns clinic artifacts into a source ledger, ambiguity queue, approved tenant policy bundle, regression evidence, controlled deployment, observable outcomes, and human-gated improvements. The recommendation is to fund source-to-policy and policy-to-test first, not a broad no-code builder or autonomous production policy learning.")
    add_table(doc, ["Decision", "What ships now", "What remains explicitly outside the proof"], [
        ("Phase 1", "Inspectable scheduling agent, architecture, scenario catalog, iteration evidence", "Telephony, real EHR/PMS, production identity, real patient data"),
        ("Phase 2", "Designed FieldFlow lifecycle, generalization boundary, experiment contracts", "Multi-tenant backend, deployment control plane, PHI analytics, autonomous learning"),
        ("P0 recommendation", "Policy compiler + generated critical scenarios", "Module marketplace, self-serve visual builder, ungated self-modification"),
    ], widths=[1.05, 2.6, 3.0], font_size=8.2)

    # Candidate method
    add_heading(doc, "Problem-Solving and Innovation Method", 1, "Candidate operating system")
    add_body(doc, "My Systemic Design Method—Observe, Frame, Build, Validate—governs how I solve a bounded product problem. My Four-Stage Innovation Method—Sense, Reason, Deploy, Improve—governs how evidence moves from signal to action and returns as learning. The first method operates within every lifecycle stage; the second spans the full field-to-production loop.")
    add_figure(doc, METHOD, "Figure 1. Dual-method operating model: problem solving inside a continuous evidence-to-outcome loop.", width=6.75)
    add_table(doc, ["Position", "Systemic Design", "Innovation", "Riverbend proof"], [
        ("01", "Observe", "Sense", "Read people, workflow, rules, dense tables, source quality, and caller signal before selecting technology"),
        ("02", "Frame", "Reason", "Define the measurable promise; separate prompt, knowledge, policy, tools, safety, identity, and assumptions"),
        ("03", "Build", "Deploy", "Connect synthetic data, state machine, typed tools, interface, trace, tests, and human handoff"),
        ("04", "Validate", "Improve", "Run negative guarantees, record defects, replay regression, propose human-gated changes"),
    ], widths=[0.45, 1.1, 1.0, 3.95], font_size=8)
    add_callout(doc, "Method guardrail", "Useful over impressive. Evidence stays in the interface. Model, data, policy, privacy, and handoff form one product system.", "attention")

    # Context and users
    add_heading(doc, "Context, Users, and Product Promise", 1, "Observe + Sense")
    add_body(doc, "Confido positions its healthcare agents as administrative operators that complete work inside EHR/PMS workflows and deliberately hand off clinical judgment.[1] The role emphasizes an FDE-first platform lifecycle—build, test, deploy, monitor, evaluate, improve, and scale—rather than a platform-first abstraction detached from clinic reality.[2]")
    add_heading(doc, "Primary Users", 2)
    add_table(doc, ["User", "Job to be done", "Success signal", "Primary risk"], [
        ("Patient or caller", "Complete an in-scope administrative task or reach the right human", "Correct, comprehensible outcome with minimal repetition", "Privacy, wrong booking, medical-advice boundary"),
        ("Forward Deployed Engineer", "Translate clinic truth into a reliable live workflow and maintain it", "Less active setup and diagnosis time without lower correctness", "Silent assumption, opaque failure, repeated custom work"),
        ("Clinic owner / operator", "Approve local policy and recover from change safely", "Source-linked review, bounded rollout, reversible release", "Automation bias, local nuance lost"),
        ("Platform engineer / quality", "Maintain shared contracts and improve cross-clinic reliability", "Reusable primitives with stable interfaces and regression evidence", "False abstraction, PHI overcollection"),
    ], widths=[1.05, 2.15, 1.75, 1.65], font_size=7.7)
    add_heading(doc, "Measurable Product Promise", 2)
    add_callout(doc, "Promise", "Complete in-scope scheduling work correctly or transfer appropriately—without disclosing appointment detail before verification, overriding clinic policy, writing before explicit confirmation, or crossing into clinical judgment.", "ink")
    add_heading(doc, "Scope", 2)
    add_table(doc, ["In scope", "Out of scope"], [
        ("Book new and follow-up visits; reschedule, cancel, confirm; hours, locations, parking; front-desk and nurse-line transfers", "Diagnosis, treatment, triage, medication or prescription work, billing resolution, annual physical configuration, identity proofing for production"),
        ("Synthetic mock tools and deterministic policy", "Live EHR/PMS, real patient data, production authorization, payment, real telephony"),
        ("Browser voice as progressive enhancement", "Voice biometric identity or a claim of full accessibility coverage"),
    ], widths=[3.2, 3.2], font_size=8)

    # Source truth and assumptions
    add_heading(doc, "Source Truth and Assumption Ledger", 1, "Observe → Frame")
    add_body(doc, "The source files were treated as evidence, not trusted configuration. High-resolution inspection corrected patient fields that small-preview OCR had misread. The provider matrix retained a real contradiction, the annual-physical workflow remained unresolved, and appointment records were clearly labeled as candidate-created fixtures. This is the behavior FieldFlow should make routine.")
    add_table(doc, ["Item", "Observed source condition", "Demo disposition", "Production owner"], [
        ("Whitfield Thursday", "Workday and location cells conflict", "No Thursday Whitfield slot; alternate-location interpretation withheld", "Clinic owner + FDE"),
        ("Annual physical", "Explicitly under consideration; not implemented", "Front-desk transfer with POLICY_TBD", "Clinic owner"),
        ("Policy number", "No policy-number values provided", "Ask only whether card/number is available; never invent or display a number", "Integration + clinic"),
        ("Existing appointments", "Three patients marked as having an upcoming appointment, without details", "Three synthetic records created and labeled as assumptions", "FDE with clinic approval"),
        ("Reference date", "Three-year and age logic depend on time", "Pinned to 2026-08-26", "Policy version owner"),
    ], widths=[1.1, 2.1, 2.25, 1.15], font_size=7.6)
    add_callout(doc, "Product principle", "An unresolved assumption is not missing documentation. It is a release artifact with an owner, severity, and disposition.", "attention")

    # Phase 1 requirements
    add_heading(doc, "Phase 1 Product Requirements", 1, "Frame + Reason")
    requirements = [
        ("P1-FR-01", "Book follow-up", "Classify within three years; 15 minutes; assigned provider or approved paired clinician; earliest eligible slot"),
        ("P1-FR-02", "Book new patient", "New or more than three years; 30 minutes; active coverage and card/policy available"),
        ("P1-FR-03", "Confirm", "Verify identity before revealing appointment; read-only confirmation"),
        ("P1-FR-04", "Cancel", "Verify, retrieve, summarize exact appointment, require explicit yes, then cancel one record"),
        ("P1-FR-05", "Reschedule", "Verify, retrieve, re-evaluate eligibility, offer earliest eligible slot, require explicit yes, mutate atomically"),
        ("P1-FR-06", "FAQ", "Answer only approved hours, lunch closure, locations, and parking"),
        ("P1-FR-07", "Human transfer", "Front desk for person request, billing, unsupported work, annual physical, or approved exceptions"),
        ("P1-FR-08", "Clinical transfer", "No medical advice; nurse line for urgent/clinical judgment; 911 instruction for possible emergency"),
        ("P1-NFR-01", "Privacy", "Phone may retrieve a candidate; DOB verifies before protected reads or writes; never disambiguate by name alone"),
        ("P1-NFR-02", "Accessibility", "Text equivalent, slower/repeat path, corrections, keyboard use, visible focus, human path without penalty"),
        ("P1-NFR-03", "Observability", "Emit redacted intent, verification, rule, tool, result, response, and escalation events"),
        ("P1-NFR-04", "Mutation safety", "No write before explicit confirmation; return structured receipt or reason-coded failure"),
    ]
    add_table(doc, ["ID", "Capability", "Acceptance contract"], requirements, widths=[0.8, 1.35, 4.25], font_size=7.8)

    # Workflows
    add_heading(doc, "Conversation and Tool Workflows", 1, "Build + Deploy")
    add_heading(doc, "Booking Sequence", 2)
    add_table(doc, ["Step", "Agent action", "Tool or rule evidence", "Failure behavior"], [
        ("1", "Interpret request; preempt emergency/clinical language", "RB-SAFE-01 / RB-SAFE-02", "911 instruction or nurse-line transfer"),
        ("2", "Request phone; retrieve candidate only", "lookupPatient", "No match → front desk; duplicate name is never enough"),
        ("3", "Verify DOB before protected information", "verifyIdentity", "Mismatch → no appointment disclosure or mutation"),
        ("4", "Confirm card/policy availability", "RB-COV-01 / RB-COV-02", "Inactive or unavailable → deny before availability"),
        ("5", "Evaluate age, discharge, visit type, provider, duration", "evaluateBookingEligibility + reason codes", "Transfer or stop on local-policy denial"),
        ("6", "Search eligible slots and sort chronologically", "searchAvailability + RB-SLOT-01", "No slot → front desk or follow clinic fallback"),
        ("7", "Repeat exact provider, location, date, time", "pending action state", "No/unclear response → no mutation"),
        ("8", "Book after explicit yes; issue receipt", "bookAppointment", "Structured failure; no false success"),
    ], widths=[0.38, 2.05, 2.1, 1.85], font_size=7.5)
    add_heading(doc, "Failure and Transfer Contract", 2)
    add_table(doc, ["Condition", "Outcome", "Forbidden behavior"], [
        ("No patient", "Front-desk path; no protected record detail", "Do not guess or create a patient"),
        ("No appointment", "Explain no upcoming record is available; front desk if needed", "Do not call cancel/reschedule"),
        ("Minor, inactive coverage, missing policy, discharged", "Reason-coded denial or front-desk transfer", "Do not search availability or write"),
        ("Different follow-up provider", "Office approval / front desk", "Do not self-book the requested provider"),
        ("Urgent clinical concern", "Nurse line; no medical advice", "Do not continue scheduling"),
        ("Possible emergency", "Tell caller to hang up and dial 911; nurse-line support only after instruction", "Do not look up patient first"),
        ("Unsupported request", "One bounded repair, then front desk", "Do not invent capability"),
    ], widths=[2.0, 2.25, 2.15], font_size=7.8)

    # Architecture
    add_heading(doc, "Runtime Architecture", 1, "One coherent system")
    add_body(doc, "The orchestration model interprets and communicates; it does not own mutable patient data or clinic authority. Safety and identity gates execute before protected work. The policy layer returns an allow/deny/transfer decision and reason codes before availability. Typed tools mediate patient, coverage, schedule, slot, appointment, and transfer actions. A tenant adapter isolates vendor behavior. The trace records enough structured evidence for replay without treating unrestricted transcripts as the default analytical substrate.")
    add_figure(doc, ARCH, "Figure 2. Riverbend runtime: agentic language, deterministic authority, typed tools, and evidence-bearing traces.", width=7.0)
    add_table(doc, ["Layer", "Owns", "Must not own"], [
        ("Prompt", "Scope, tone, safety order, turn-taking, verification sequence, confirmation discipline", "Patient records, hidden eligibility logic, provider schedules"),
        ("Knowledge", "Approved read-only hours, locations, parking, service boundaries", "Mutable appointments, unresolved policy"),
        ("Policy", "Age, coverage, discharge, visit type, duration, provider, earliest-slot order", "Conversational phrasing or medical advice"),
        ("Tools", "Validated reads, availability, mutations, transfers, receipts", "Free-form business judgment"),
        ("Adapter", "Vendor field mapping, auth, idempotency, structured vendor errors", "Clinic rule semantics"),
        ("Trace/evaluation", "Redacted events, assertions, replay, release evidence", "Unrestricted PHI analytics"),
    ], widths=[0.9, 3.0, 2.5], font_size=7.8)
    add_body(doc, "FHIR separates an Appointment from the Schedule and Slot resources used to represent availability, and it notes that additional business rules may be required before a booking is allowed.[3] The prototype uses that conceptual boundary without claiming FHIR conformance: eligibility is decided above availability, and tenant adapters can later map the normalized contracts to a specific EHR/PMS.")

    # Data and rules
    add_heading(doc, "Policy, Data, and Interoperability Contracts", 1, "Reason")
    add_table(doc, ["Entity", "Minimum contract", "Privacy / authority boundary"], [
        ("Patient", "tenant ID, external ID, name, DOB, phone, assigned provider, discharge flag", "Candidate lookup is not verification"),
        ("Coverage", "payer, status, policy/card-available state", "No invented policy number; status must precede availability"),
        ("Provider", "ID, role, location, pairing, work pattern", "Local exceptions remain tenant policy"),
        ("Schedule", "provider, location, service, planning horizon", "Availability is not eligibility"),
        ("Slot", "start/end, status, provider, location", "Offer only after approved decision"),
        ("Appointment", "patient, provider, slot/time, type, duration, status, version", "Read/write only after verification; writes after confirmation"),
        ("Outcome", "status, reason code, message class, receipt, correlation ID", "Do not infer success from prose"),
    ], widths=[0.85, 3.35, 2.2], font_size=7.7)
    add_heading(doc, "Core Rule Order", 2)
    add_bullets(doc, [
        "Safety preemption before intent completion.",
        "Candidate lookup before verification; verification before protected read or write.",
        "Age, discharge, coverage, and card/policy availability before provider or slot search.",
        "Visit classification and duration before eligible-slot filtering.",
        "Assigned-provider continuity and Dr. Crane rules before availability.",
        "Chronological earliest eligible offer before alternatives.",
        "Exact action summary and explicit confirmation before mutation.",
    ])

    # Safety privacy accessibility
    add_heading(doc, "Safety, Privacy, Security, and Accessibility", 1, "Human first")
    add_body(doc, "HHS states that covered entities should use reasonable safeguards for protected health information and verify the identity of an unknown individual, while not prescribing one universal verification method.[4] It also notes that audio, recordings, and electronic communications can create Security Rule obligations. For Riverbend, phone matching retrieves a candidate and DOB verifies the person before appointment disclosure or mutation. Production deployment would require clinic-specific identity risk decisions, encryption, access controls, retention, BAA/vendor review, and incident response.")
    add_body(doc, "HHS and DOJ guidance on effective communication supports offering alternative communication methods and accessible scheduling interactions.[5] The prototype therefore treats voice as optional: every path can be completed through text or scripted scenarios; focus is visible; controls are labeled; motion can be disabled; and a human path is always available. An English-only browser demo is not represented as complete language or disability coverage.")
    add_table(doc, ["Control", "Prototype evidence", "Production extension"], [
        ("Identity", "Phone candidate + DOB gate", "Configurable attributes, accessible fallback, rate limit, audit"),
        ("Least privilege", "Typed tool boundaries and forbidden-call tests", "Scoped credentials, RBAC, service identities"),
        ("PHI-aware observability", "Structured synthetic trace", "Redaction, retention, access review, secure export"),
        ("Human escalation", "Front desk, nurse line, 911 instruction", "Warm-transfer state, fallback verification, queue health"),
        ("Accessibility", "Text path, keyboard, focus, reduced motion", "Language routing, relay/support pathways, user research"),
        ("Release safety", "Critical scenario suite", "Environment promotion, canary, rollback, incident playbook"),
    ], widths=[1.05, 2.3, 3.05], font_size=7.7)

    # Tests
    add_heading(doc, "Evaluation Strategy and Actual Results", 1, "Validate")
    add_body(doc, "A believable final sentence is not sufficient evidence. The test strategy inspects the outcome and the path beneath it: reason code, rule IDs, forbidden tools, required tools, state mutation, transfer destination, and source-grounded knowledge. Critical failures block release.")
    add_table(doc, ["Validation", "Actual result", "Interpretation"], [
        ("Automated assertions", "18 / 18 passed", "Policy boundaries, privacy, safety, mutation confirmation, slot order, and catalog replay"),
        ("Scenario catalog", "24 / 24 replayed successfully", "Happy paths, denials, identity, FAQs, transfers, urgent and emergency cases"),
        ("TypeScript", "No errors", "Typed client policy and interaction model are internally consistent"),
        ("Production build", "Succeeded", "Deployable static artifact; one non-blocking large-chunk optimization warning remains"),
        ("Visual review", "6 representative routes captured", "Field-notebook direction confirmed; failed first-generation images were detected and replaced"),
    ], widths=[1.35, 1.35, 3.95], font_size=8)
    add_heading(doc, "Critical Release Contract", 2)
    add_callout(doc, "Zero tolerance", "No release may book a minor, discharged patient, inactive coverage, missing policy/card, or unapproved provider change; reveal protected appointment data after failed verification; call a write before explicit confirmation; or continue scheduling through urgent clinical or emergency language.", "attention")

    # Iteration
    add_heading(doc, "High-Level Iteration Log", 1, "What changed and why")
    add_table(doc, ["Iteration", "Change", "Trigger", "Product impact"], [
        ("0", "Re-read dense tables at high resolution", "OCR and preview errors", "Source quality precedes automation"),
        ("1", "Rejected prompt-only rules", "Hard to test, diff, or guarantee", "Versioned deterministic policy"),
        ("2", "Strengthened identity gate", "Duplicate James Porter records", "DOB gates disclosure and mutation"),
        ("3", "Made safety preemptive", "Clinical language can arrive at any turn", "Scheduling stops before transfer/911 instruction"),
        ("4", "Moved availability after eligibility", "A free slot does not imply authorization", "Forbidden searches for denied patients"),
        ("5", "Created ambiguity queue", "Conflicting/missing source truth", "Conservative labeled assumptions"),
        ("6", "Isolated appointment assumptions", "Brief leaves details to candidate", "Reproducible confirm/cancel/reschedule"),
        ("7", "Expanded assertions beyond prose", "Fluent text can hide a bad workflow", "Tool, reason, state, and transfer evidence"),
        ("8", "Made voice optional", "Telephony deprioritized; browser support varies", "Reliable evaluator and accessibility path"),
        ("9", "Mapped friction into FieldFlow", "Phase 2 asks for generalization", "Productization begins with observed FDE work"),
        ("10", "Replaced failed visual assets", "Full-page visual review", "Defect recorded, fixed, and revalidated"),
    ], widths=[0.55, 1.65, 2.0, 2.2], font_size=7.4)

    # Phase2 thesis
    add_heading(doc, "Phase 2: FieldFlow Platform Strategy", 1, "From field learning to platform leverage")
    add_body(doc, "The hard work in Riverbend was a repeatable control problem: preserve source truth, clarify ambiguity, compile authority into stable layers, simulate real failure modes, approve the exact version, deploy reversibly, observe bounded outcomes, and improve without ungoverned policy drift. FieldFlow productizes that work for FDEs while preserving clinic-specific truth.")
    add_figure(doc, LIFECYCLE, "Figure 3. FieldFlow lifecycle: every handoff creates an inspectable artifact and a release decision.", width=7.0)
    add_table(doc, ["Stage", "Question", "Artifact", "Critical gate"], [
        ("Ingest", "What did the clinic provide?", "Source ledger", "Unreadable critical source blocks compilation"),
        ("Clarify", "What is missing or conflicting?", "Ambiguity queue", "Critical item requires owner and disposition"),
        ("Compile", "What belongs in each layer?", "Tenant policy bundle", "Human approval for critical policy"),
        ("Simulate", "Can it survive real failures?", "Regression matrix", "All critical scenarios pass"),
        ("Approve", "Is local truth represented?", "Release candidate", "Named owner and rollback"),
        ("Deploy", "How is risk bounded?", "Deployment record", "Adapter health and canary scope"),
        ("Observe", "Did it resolve appropriately?", "Quality dashboard", "Threshold alert and incident path"),
        ("Improve", "What deserves change?", "Change proposal", "Human approval and regression replay"),
    ], widths=[0.82, 2.0, 1.65, 1.95], font_size=7.6)

    # Generalization
    add_heading(doc, "Generalization Boundary", 1, "Contracts, not copy-paste")
    add_body(doc, "A reusable primitive earns platform status when it is cross-tenant, versionable, permissionable, observable, and testable. A rule remains tenant configuration when its meaning depends on local policy, staffing, service scope, location, or unresolved clinic judgment.")
    add_table(doc, ["Generalize", "Why it compounds", "Keep local or configurable", "Why"], [
        ("Identity workflow", "Protects reads and writes", "Required attributes and fallback", "Accessible/risk practices vary"),
        ("Policy schema + compiler", "Makes authority diffable", "Thresholds, durations, pairings", "Clinic policy varies"),
        ("Typed tool registry", "Standardizes authorization and receipts", "Vendor adapter and permissions", "EHR/PMS behavior varies"),
        ("Scenario families", "Reuses safety and mutation patterns", "Fixtures and local expected outcomes", "Truth follows the tenant bundle"),
        ("Trace, replay, release, rollback", "Reduces diagnosis and migration risk", "Retention and alert thresholds", "Privacy and volume vary"),
        ("Module registry", "Surfaces field-proven patterns", "Acceptance and overrides", "Every destination needs local approval"),
        ("Ambiguity workflow", "Prevents silent assumptions", "Owner and disposition", "Only the clinic can settle some questions"),
    ], widths=[1.25, 1.65, 1.9, 1.6], font_size=7.4)
    add_callout(doc, "Riverbend stays local", "Three-year threshold, Dr. Crane Thursday, provider pairings, lunch closure, annual-physical status, and unresolved Whitfield interpretation never become defaults merely because they worked here.", "ink")

    # Human gated learning
    add_heading(doc, "Human-Gated Improvement", 1, "Improve without policy drift")
    add_body(doc, "Runtime learning should detect recurring patterns, attach evidence, identify the likely layer, draft a source-linked change, and replay regression. It should not rewrite eligibility or escalation behavior from live calls. Approved tenant bundles remain immutable; a suggestion becomes a release only after authorized human review, critical tests, canary scope, and rollback readiness.")
    add_table(doc, ["Signal", "Platform action", "Human decision", "Promotion evidence"], [
        ("Repeated unsupported intent", "Cluster redacted events and propose a new intent or transfer refinement", "FDE/clinic approves scope", "Scenario, trace, no-regression suite"),
        ("Policy denial confusion", "Propose clearer response or source clarification", "Clinic confirms wording and rule", "Comprehension test + unchanged authority"),
        ("Adapter failure", "Identify tool/vendor layer and suggest fallback", "Engineering approves technical change", "Contract tests, canary, rollback"),
        ("Cross-clinic pattern", "Suggest reusable module with applicability and provenance", "Platform review + destination acceptance", "Repeated evidence and zero inherited critical errors"),
    ], widths=[1.25, 2.1, 1.6, 1.55], font_size=7.6)

    # Experiments
    add_heading(doc, "Prioritized Experiment Portfolio", 1, "Want the answer before wanting the feature")
    add_body(doc, "Priority reflects dependency order, safety, and the value of the answer. Each experiment has a smallest useful test, a metric contract, and a rule to scale, narrow, repeat, or stop. P0 validates the central source-to-test thesis before broader module or runtime-learning investment.")
    add_table(doc, ["ID / Priority", "Hypothesis", "Smallest useful test", "Metric", "Want answer", "Decision rule"], [
        ("EXP-01 / P0", "Compiler cuts setup ≥30% without critical regression", "Two FDEs; two packets; counterbalanced current vs FieldFlow", "Median active time + critical pass rate", "10/10", "Scale at ≥30% and non-inferior correctness; stop on critical regression"),
        ("EXP-02 / P0", "Generated scenarios find ≥25% more critical defects per reviewer hour", "Seed identical policy/tool defects; generated vs manual QA", "Critical defects/hour + false positives", "10/10", "Scale at ≥25% lift and <15% false positives"),
        ("EXP-03 / P1", "Source-linked traces cut diagnosis time ≥40%", "Randomize matched failures: transcript vs trace", "Time + correct layer", "9/10", "Scale with no excess PHI exposure"),
        ("EXP-04 / P1", "Ambiguity queue catches all planted critical conflicts", "Three packets with contradictory/missing fields", "Critical conflict at approval + review time", "9/10", "Scale at zero misses and ≤15% extra review"),
        ("EXP-05 / P2", "Modules improve second-clinic velocity", "Configure second clinic with mandatory source diff", "Acceptance + overrides + critical regressions", "7/10", "Scale at ≥50% acceptance, zero inherited critical errors"),
        ("EXP-06 / P2", "Human-gated suggestions improve replay", "Shadow-mode failure clustering and FDE review", "Acceptance + replay gain + review time", "6/10", "Expand at ≥30% acceptance, zero approval bypass"),
    ], widths=[0.72, 1.55, 1.55, 1.15, 0.62, 1.45], font_size=6.8)
    add_callout(doc, "Roadmap decision", "Fund the policy compiler and generated critical scenarios first. Defer a broad visual builder, module propagation, and runtime learning until P0 proves FDE leverage without lower correctness.", "verified")

    # Metrics roadmap
    add_heading(doc, "Metrics, Rollout, and Roadmap", 1, "Deploy + Improve")
    add_heading(doc, "North Star", 2)
    add_callout(doc, "Clinically bounded administrative resolution", "The share of in-scope requests that end in a correct, policy-compliant administrative outcome or appropriate transfer, with no safety or privacy violation.", "ink")
    add_table(doc, ["Dimension", "Primary metric", "Guardrail"], [
        ("FDE leverage", "Active FDE hours per launch; time to approved release", "Critical pass rate must not decline"),
        ("Patient outcome", "Bounded resolution; repeat/correction rate; appropriate transfer", "Zero known medical-advice completions"),
        ("Quality", "Critical regression rate; policy/tool error rate", "Release blocked by critical failure"),
        ("Operations", "Diagnosis time; recovery time; rollback success", "Trace is role-scoped and PHI-aware"),
        ("Compounding", "Module acceptance; second-clinic setup time", "Zero inherited critical rule errors"),
        ("Trust", "Review turnaround; ambiguity age; rollback requests", "Critical ambiguity requires disposition"),
    ], widths=[1.0, 3.1, 2.3], font_size=7.8)
    add_heading(doc, "Sequenced Roadmap", 2)
    add_table(doc, ["Horizon", "Build", "Learn", "Do not add yet"], [
        ("Now", "Source ledger, ambiguity queue, typed policy bundle, critical scenarios", "Time, correctness, defect yield, review burden", "Broad drag-and-drop builder"),
        ("Next", "Trace replay, source-linked diffs, release gates, rollback", "Diagnosis speed, reviewer comprehension, migration safety", "Cross-clinic auto-acceptance"),
        ("Later", "Controlled module suggestions, shadow-mode improvement proposals", "Second-clinic leverage and accepted suggestion quality", "Ungated production self-modification"),
    ], widths=[0.8, 2.25, 2.15, 1.2], font_size=7.8)

    # Risks
    add_heading(doc, "Risks and Decision Controls", 1, "Govern")
    add_body(doc, "NIST’s AI Risk Management Framework emphasizes governance, mapping, measurement, and management throughout the lifecycle.[6] Riverbend and FieldFlow translate that into named policy ownership, explicit users and harms, critical evaluation contracts, bounded rollout, rollback, escalation, and incident evidence.")
    add_table(doc, ["Risk", "Failure mode", "Control", "Owner"], [
        ("False abstraction", "Local rule becomes platform default", "Provenance, local approval, overrides, cross-clinic evidence threshold", "Platform PM + FDE"),
        ("Automation bias", "Reviewer accepts a confident but wrong extraction", "Source excerpt, confidence, explicit critical-field review", "FDE + clinic owner"),
        ("Policy drift", "Runtime changes authority", "Immutable approved bundles; draft-only suggestions", "Platform engineering"),
        ("PHI overcollection", "Unnecessary transcript/trace exposure", "Structured redacted events, retention, RBAC", "Security/privacy"),
        ("Test theater", "Passing prose hides unsafe tools", "Forbidden-call, reason-code, and mutation-state assertions", "Quality owner"),
        ("Platform burden", "Workflow adds time without lowering defects", "Counterbalanced pilot and explicit stop rules", "Platform PM"),
    ], widths=[1.0, 1.65, 2.85, 0.9], font_size=7.5)

    # Candidate fit
    add_heading(doc, "Why This Work Reflects the Role", 1, "Candidate operating fit")
    add_body(doc, "This submission demonstrates the operating pattern the role describes: embed in a messy healthcare workflow, reason across model, data, tools, privacy, UI, and handoff, work in the codebase, ship an inspectable artifact quickly, define lifecycle contracts, test negative guarantees, and convert field friction into platform primitives and falsifiable experiments.")
    add_table(doc, ["Role signal", "Evidence in this assignment", "Candidate background"], [
        ("Agentic AI end to end", "State machine, prompt, policy, tools, voice, trace, scenarios, release concepts", "AI-agent platforms, evaluation frameworks, API-first production workflows"),
        ("Healthcare operations", "Scheduling, coverage, identity, EHR/PMS boundary, clinical handoff", "Digital health, healthcare workflows, clinical trials, radiomics, precision health"),
        ("Hands-on builder PM", "React/TypeScript prototype, tests, architecture, open-source repository", "Product engineering, Python/TypeScript, cloud, CI/CD, AI-native tooling"),
        ("FDE-first platform", "FieldFlow lifecycle, ambiguity ownership, source-linked review, adapters", "Translates operator ambiguity into contracts and measurable systems"),
        ("Disciplined experiments", "Six prioritized falsifiable bets with stop/scale rules", "Hypothesis → smallest test → honest read operating style"),
        ("NYC and field proximity", "Designed for clinic/FDE co-ownership and travel-based discovery", "NYC-based, U.S. citizen, available for onsite work and domestic travel"),
    ], widths=[1.35, 2.9, 2.15], font_size=7.6)
    add_callout(doc, "Leadership stance", "Build beside the domain expert. Make the decision inspectable. Scale only what the evidence earns.", "attention")

    # Deliverables
    add_heading(doc, "Deliverable Index", 1, "Submission map")
    add_table(doc, ["Assignment item", "Deliverable", "Where to inspect"], [
        ("The agent", "Interactive voice/text scheduling workspace", "Prototype `/agent`; `client/src/lib/agentEngine.ts`; `agent/system-prompt.md`"),
        ("Scenarios", "24-case catalog and interactive runner", "Prototype `/tests`; `agent/scenarios.json`; automated tests"),
        ("Architecture write-up", "Prompt/knowledge/policy/tool/runtime design", "Prototype `/architecture`; `docs/architecture.md`"),
        ("Iteration log", "Material changes, triggers, impact, actual validation", "`docs/iteration-log.md`"),
        ("Part 2 short document", "FieldFlow strategy", "`docs/phase-2-platform-strategy.md`"),
        ("Designed prototype", "Eight-stage lifecycle and boundary board", "Prototype `/platform`"),
        ("Experiments", "P0–P2 portfolio and honest-read contracts", "Prototype `/experiments`; `docs/experiment-portfolio.md`"),
        ("Video support", "Interactive PowerPoint and timed script", "Presentation and `docs/designed-prototype-guide.md`"),
        ("Open source", "Documented repository, tests, license, CI", "Riverbend Gastroenterology Scheduling repository"),
    ], widths=[1.55, 2.2, 2.65], font_size=7.7)

    # Appendix scenarios in landscape
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=True)
    add_heading(doc, "Appendix A — Complete Scenario Catalog", 1, "24 replayable proofs")
    scenario_data = json.loads(SCENARIOS.read_text())
    scenario_rows = []
    for item in scenario_data["scenarios"]:
        expected = item.get("expected", {})
        evidence_parts = []
        if expected.get("reasonCode"):
            evidence_parts.append(expected["reasonCode"])
        if expected.get("requiredRules"):
            evidence_parts.append("Rules: " + ", ".join(expected["requiredRules"]))
        if expected.get("requiredTools"):
            evidence_parts.append("Tools: " + ", ".join(expected["requiredTools"]))
        if expected.get("forbiddenTools"):
            evidence_parts.append("Forbidden: " + ", ".join(expected["forbiddenTools"]))
        scenario_rows.append((
            item["id"], item["title"], item["criticality"].upper(), item["intent"], expected.get("outcome", "—"), "; ".join(evidence_parts) or "Outcome assertion",
        ))
    add_table(doc, ["ID", "Scenario", "Risk", "Intent", "Expected outcome", "Evidence assertion"], scenario_rows, widths=[1.45, 2.35, 0.72, 1.25, 1.35, 2.65], font_size=6.5)

    # Appendix trace and requirements
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=False)
    add_heading(doc, "Appendix B — Trace and Release Evidence", 1, "Inspectability contract")
    add_table(doc, ["Event field", "Purpose", "Example / constraint"], [
        ("correlationId", "Connect turns, tools, and receipt", "Opaque; tenant-scoped"),
        ("bundleVersion", "Reproduce policy authority", "1.0.0"),
        ("stage", "Locate the failure layer", "intent, safety, identity, policy, tool, response"),
        ("verificationState", "Prove protected operation gate", "candidate, verified, failed, not-required"),
        ("ruleIds / reasonCode", "Tie outcome to approved policy", "RB-AGE-01 / UNDER_18"),
        ("tool / requestClass", "Prove allowed action sequence", "searchAvailability; no raw credentials"),
        ("toolResultClass", "Distinguish success from structured failure", "receipt, no-match, denied, vendor-error"),
        ("responseClass", "Evaluate behavior without defaulting to full transcript", "offer, deny, transfer, emergency instruction"),
        ("escalation", "Verify human destination", "front desk, nurse line, 911 instruction"),
        ("privacy", "Limit analytical exposure", "Redaction, retention, role-scoped access"),
    ], widths=[1.35, 2.35, 2.7], font_size=7.8)
    add_heading(doc, "Release Checklist", 2)
    add_bullets(doc, [
        "Critical source fields reviewed; ambiguity owners and dispositions recorded.",
        "Prompt, knowledge, policy, tool, and adapter diffs approved by their accountable owners.",
        "All critical scenarios pass with required and forbidden tool evidence.",
        "Production identity, authorization, privacy, accessibility, and transfer behaviors are reviewed for the tenant.",
        "Adapter health, idempotency, canary scope, alert thresholds, and rollback owner are confirmed.",
        "North-star and guardrail baselines are captured before rollout.",
    ])

    add_heading(doc, "Appendix C — References", 1, "Primary and authoritative sources")
    refs = [
        ("Confido Health — Voice AI Receptionist for Healthcare Practices", "https://www.confido.health/home"),
        ("Confido Health — Product Manager, Platform role", "https://www.paraform.com/share/confido-health/cmqfno2if000a0cl8js9ax108"),
        ("HL7 FHIR R5 — Appointment", "https://www.hl7.org/fhir/appointment.html"),
        ("HHS OCR — HIPAA and Audio-Only Telehealth", "https://www.hhs.gov/hipaa/for-professionals/privacy/guidance/hipaa-audio-telehealth/index.html"),
        ("HHS/DOJ — Nondiscrimination in Telehealth", "https://www.hhs.gov/civil-rights/for-individuals/disability/guidance-on-nondiscrimination-in-telehealth/index.html"),
        ("NIST — AI Risk Management Framework", "https://www.nist.gov/itl/ai-risk-management-framework"),
        ("Cekura — Confido Health case study", "https://www.cekura.ai/case-study/confido-health"),
        ("Argonaut Scheduling Implementation Guide", "https://fhir.org/guides/argonaut/scheduling/patient-scheduling.html"),
        ("Eangelica Aton — Innovation and Systemic Design Methods", "https://2026.eangelica.com/#method"),
    ]
    for index, (title, url) in enumerate(refs, start=1):
        add_reference(doc, index, title, url)

    add_heading(doc, "Appendix D — Prototype and Repository Notes", 1, "Reproducibility")
    add_body(doc, "The repository is designed for evaluator inspection and open-source reuse. The runnable prototype is a React/TypeScript static application. The deterministic engine, synthetic fixtures, policy bundle, knowledge base, scenario catalog, tests, strategy documents, diagrams, and presentation source are versioned together. Visual assets are hosted separately for deploy reliability; diagram source remains in the repository.")
    add_table(doc, ["Command", "Purpose"], [
        ("pnpm install", "Install locked dependencies"),
        ("pnpm dev", "Run the local evaluator experience"),
        ("pnpm test", "Run 18 deterministic assertions, including all 24 scenario replays"),
        ("pnpm check", "Run TypeScript validation"),
        ("pnpm build", "Create the production artifact"),
    ], widths=[1.55, 4.85], font_size=8)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Created {OUT}")


if __name__ == "__main__":
    build_document()
